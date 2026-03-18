import os
import re
import shutil
import asyncio
import zipfile
import html
import gc
import threading
from flask import Flask
from docx import Document
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, filters, ContextTypes
from telegram.error import RetryAfter

# ==========================================
# CONFIGURATION
# ==========================================
TOKEN ="7954112414:AAEXxK5PcckCiZChIWUXHsQSRD2XfiasW-Q"
GROUP_ID = -1003745983576
DEFAULT_DOCX_CHUNK = 50
DEFAULT_EPUB_CHUNK = 500

# Global State
document_queue = None  # Will be initialized in async context
user_chunk_sizes = {}
pending_uploads = {}

# ==========================================
# LOGIC: CHAPTER DETECTION & SPLITTING
# ==========================================
def split_text_based_logic(input_path, output_dir, chunk_size, output_format, is_txt_file=False):
    if not os.path.exists(output_dir): 
        os.makedirs(output_dir)
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    generated_files = []

    # Enhanced regex for complex titles: "अध्याय 2, राक्षसी सम्राट" or "Chapter 1: Begin"
    pattern_num = re.compile(
        r"(?:vol(?:ume)?\s*\d+\s*)?"
        r"(?:chapter|ch|c|अध्याय|चैप्टर|सी|page|पृष्ठ|#)?\s*"
        r"(\d+)"
        r"(?:[:\s,.-]|$)",
        re.IGNORECASE
    )

    def save_chunk(lines, start, end, chunk_num):
        ext = ".txt" if output_format == "txt" else ".docx"
        if end == "End":
            suffix = f"Chapters_{start}_to_End"
        else:
            suffix = f"Chapters_{start}_to_{end}"
        part_name = f"{suffix}_{base_name}_Part{chunk_num}{ext}"
        part_path = os.path.join(output_dir, part_name)

        try:
            if output_format == "txt":
                with open(part_path, "w", encoding="utf-8", errors='ignore') as f:
                    f.write("\n\n".join(lines))
            else:
                new_doc = Document()
                for line in lines:
                    if line.strip(): 
                        new_doc.add_paragraph(line.strip())
                new_doc.save(part_path)
            
            generated_files.append(part_path)
            return True
        except Exception as e:
            print(f"❌ Error saving {part_name}: {e}")
            return False

    # Load Lines with better error handling
    lines = []
    try:
        if is_txt_file:
            with open(input_path, "r", encoding="utf-8", errors='ignore') as f:
                lines = [line.strip() for line in f if line.strip()]
        else:
            from docx.opc.exceptions import PackageNotFoundError
            try:
                doc = Document(input_path)
                lines = [re.sub(r'\x00', '', p.text).strip() for p in doc.paragraphs if p.text.strip()]
                del doc
            except PackageNotFoundError:
                print("❌ Error: Invalid DOCX file")
                return []
    except Exception as e:
        print(f"❌ Error reading file: {e}")
        return []

    if not lines:
        print("❌ No content found")
        return []

    # Chunking Loop
    current_chunk = []
    current_start_chapter = None
    last_detected_chapter = None
    chunk_count = 1
    total_chapters = 0

    for line in lines:
        match = pattern_num.search(line)
        detected_num = int(match.group(1)) if match else None

        if detected_num is not None:
            total_chapters += 1
            if current_start_chapter is None:
                current_start_chapter = detected_num
            
            # Boundary check: If we hit the next chunk threshold
            if detected_num >= (current_start_chapter + chunk_size):
                if save_chunk(current_chunk, current_start_chapter, last_detected_chapter or (detected_num-1), chunk_count):
                    chunk_count += 1
                current_chunk = []
                current_start_chapter = detected_num
            
            last_detected_chapter = detected_num

        current_chunk.append(line)

    # Save final chunk
    if current_chunk:
        save_chunk(current_chunk, current_start_chapter or 1, last_detected_chapter or "End", chunk_count)

    print(f"🎉 Splitting complete! Found {total_chapters} chapters, created {len(generated_files)} files")
    gc.collect()
    return generated_files

# ==========================================
# EPUB LOGIC
# ==========================================
def split_epub_logic(input_path, output_dir, chunk_size, output_format):
    if not os.path.exists(output_dir): 
        os.makedirs(output_dir)
    base_name = re.sub(r'[^\w\-_]', '_', os.path.splitext(os.path.basename(input_path))[0])
    generated_files = []

    def save_epub_chunk(lines, count):
        ext = ".txt" if output_format == "txt" else ".docx"
        part_name = f"Part_{count}-{base_name}{ext}"
        part_path = os.path.join(output_dir, part_name)
        try:
            if output_format == "txt":
                with open(part_path, "w", encoding="utf-8") as f: 
                    f.write("\n\n".join(lines))
            else:
                new_doc = Document()
                for line in lines: 
                    new_doc.add_paragraph(line)
                new_doc.save(part_path)
            generated_files.append(part_path)
        except Exception as e:
            print(f"❌ Error saving EPUB chunk: {e}")

    try:
        with zipfile.ZipFile(input_path, 'r') as epub:
            files = sorted([f for f in epub.namelist() if f.lower().endswith(('.html', '.xhtml'))])
            buffer, ch_in_chunk, chunk_idx = [], 0, 1
            for f_name in files:
                content = epub.read(f_name).decode('utf-8', errors='ignore')
                clean_lines = [html.unescape(re.sub(r'<[^>]+>', '', line)).strip() 
                               for line in re.split(r'</?(?:p|div|br|h[1-6])[^>]*>', content) if line.strip()]
                if clean_lines:
                    buffer.extend(clean_lines)
                    buffer.append("-" * 20)
                    ch_in_chunk += 1
                if ch_in_chunk >= chunk_size:
                    save_epub_chunk(buffer, chunk_idx)
                    buffer, ch_in_chunk, chunk_idx = [], 0, chunk_idx + 1
                    gc.collect()
            if buffer: 
                save_epub_chunk(buffer, chunk_idx)
    except Exception as e: 
        print(f"EPUB Error: {e}")
    return generated_files

# ==========================================
# ASYNC WORKER & HANDLERS
# ==========================================
async def queue_worker():
    while True:
        job = await document_queue.get()
        try:
            loop = asyncio.get_event_loop()
            format_name = "TXT" if job['format'] == "txt" else "DOCX"
            
            await job['status_msg'].edit_text(f"⚡ Processing {job['type'].upper()} file...")

            if job['type'] == 'docx':
                files = await loop.run_in_executor(None, split_text_based_logic, job['input_path'], job['output_dir'], job['chunk_size'], job['format'], False)
            elif job['type'] == 'txt':
                files = await loop.run_in_executor(None, split_text_based_logic, job['input_path'], job['output_dir'], job['chunk_size'], job['format'], True)
            else:
                files = await loop.run_in_executor(None, split_epub_logic, job['input_path'], job['output_dir'], job['chunk_size'], job['format'])

            if files:
                thread_id = None
                try:
                    topic = await job['context'].bot.create_forum_topic(chat_id=GROUP_ID, name=job['base_name'][:128])
                    thread_id = topic.message_thread_id
                    await job['status_msg'].edit_text(f"✅ Uploading {len(files)} files to group...")
                except Exception as e:
                    await job['status_msg'].edit_text(f"⚠️ Topic creation failed, sending to main chat")

                for f in files:
                    try:
                        with open(f, 'rb') as doc:
                            send_args = {
                                "chat_id": GROUP_ID,
                                "document": doc,
                                "filename": os.path.basename(f)
                            }
                            if thread_id:
                                send_args["message_thread_id"] = thread_id
                            
                            await job['context'].bot.send_document(**send_args)
                        await asyncio.sleep(0.5)  # Rate limiting
                    except RetryAfter as e:
                        print(f"⚠️ Rate limited, waiting {e.retry_after}s")
                        await asyncio.sleep(e.retry_after)
                    except Exception as e:
                        print(f"⚠️ Error sending file: {e}")
                
                await job['status_msg'].edit_text(f"✅ Completed: {job['base_name']} - {len(files)} files uploaded")
            else:
                await job['status_msg'].edit_text("❌ No chapters detected or processing failed.")
        except Exception as e:
            await job['status_msg'].edit_text(f"❌ Error: {str(e)}")
        finally:
            if os.path.exists(job['temp_dir']): 
                shutil.rmtree(job['temp_dir'])
            document_queue.task_done()

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    name = update.effective_user.first_name
    await update.message.reply_text(
        f"👋 Hello {name}!\n\n"
        f"📚 Send me a **.docx**, **.txt**, or **.epub** file to split.\n"
        f"⚙️ Use `/set 100` to change chunk size.\n"
        f"📤 Files will be sent to your group."
    )

async def set_chunk_size(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        new_size = int(context.args[0])
        user_chunk_sizes[update.effective_user.id] = new_size
        await update.message.reply_text(f"✅ Chunk size set to **{new_size}** chapters.")
    except:
        await update.message.reply_text("⚠️ Usage: `/set 100`")

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    doc = update.message.document
    file_name = doc.file_name.lower()
    msg_id = update.message.message_id

    if not (file_name.endswith('.docx') or file_name.endswith('.epub') or file_name.endswith('.txt')):
        await update.message.reply_text("❌ Only .docx, .txt, or .epub files allowed.")
        return

    pending_uploads[msg_id] = {
        'document': doc, 
        'user_id': update.effective_user.id, 
        'user_mention': update.effective_user.first_name
    }
    
    keyboard = [[InlineKeyboardButton("📄 DOCX", callback_data=f"file|docx|{msg_id}"), 
                 InlineKeyboardButton("📝 TXT", callback_data=f"file|txt|{msg_id}")]]
    await update.message.reply_text("Choose output format:", reply_markup=InlineKeyboardMarkup(keyboard))

async def button_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    _, fmt, msg_id = query.data.split("|")
    msg_id = int(msg_id)

    if msg_id not in pending_uploads: 
        await query.edit_message_text("❌ Session expired. Please upload again.")
        return
        
    job_info = pending_uploads.pop(msg_id)
    
    # Use /tmp for Render compatibility
    temp_dir = os.path.join("/tmp", f"job_{msg_id}")
    os.makedirs(os.path.join(temp_dir, "output"), exist_ok=True)
    input_path = os.path.join(temp_dir, job_info['document'].file_name)

    status = await query.edit_message_text("📥 Downloading file...")
    file_obj = await job_info['document'].get_file()
    await file_obj.download_to_drive(input_path)

    # Determine file type
    if input_path.endswith('.txt'):
        file_type = 'txt'
    elif input_path.endswith('.epub'):
        file_type = 'epub'
    else:
        file_type = 'docx'

    await document_queue.put({
        'type': file_type,
        'format': fmt, 
        'chunk_size': user_chunk_sizes.get(job_info['user_id'], DEFAULT_DOCX_CHUNK),
        'input_path': input_path, 
        'output_dir': os.path.join(temp_dir, "output"),
        'temp_dir': temp_dir, 
        'base_name': os.path.splitext(job_info['document'].file_name)[0],
        'status_msg': status, 
        'context': context, 
        'user_mention': job_info['user_mention']
    })

# ==========================================
# WEB SERVER & MAIN
# ==========================================
app_web = Flask(__name__)

@app_web.route('/')
def home(): 
    return "Bot Running"

@app_web.route('/health')
def health():
    return {"status": "healthy"}

def run_web():
    port = int(os.environ.get("PORT", 8080))
    app_web.run(host="0.0.0.0", port=port)

async def start_background_tasks(app: Application):
    global document_queue
    document_queue = asyncio.Queue()  # Initialize in async context
    asyncio.create_task(queue_worker())

def main():
    # Start web server in background
    threading.Thread(target=run_web, daemon=True).start()
    
    # Initialize Telegram bot
    app = Application.builder().token(TOKEN).post_init(start_background_tasks).build()
    
    # Add handlers
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("set", set_chunk_size))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(CallbackQueryHandler(button_callback))
    
    print("🚀 Bot Started on Render")
    app.run_polling(drop_pending_updates=True)

if __name__ == "__main__":
    main()
